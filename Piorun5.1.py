# Generator dokumentacji Krok po Kroku - Piorun 5.1
# Wersja z PySide6 (Qt for Python)

import sys
import os
import json
import base64
import tempfile
import shutil
import copy
import traceback
import hashlib
from datetime import datetime

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QListWidget, QTextEdit,
    QGroupBox, QComboBox, QFileDialog, QMessageBox, QFrame,
    QGridLayout, QSplitter, QProgressDialog, QListWidgetItem,
    QSizePolicy, QMenuBar, QMenu, QStyleFactory
)
from PySide6.QtCore import Qt, QSize, QTimer, QThread, Signal
from PySide6.QtGui import QPixmap, QFont, QPalette, QColor, QIcon, QAction
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import subprocess


class ThumbnailListWidget(QListWidget):
    """ListWidget z miniaturami obrazów i numerami kroków"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setIconSize(QSize(80, 80))
        self.setResizeMode(QListWidget.Adjust)
        self.setViewMode(QListWidget.IconMode)
        self.setMovement(QListWidget.Static)
        self.setSpacing(15)
        self.setMinimumHeight(140)
        self.setMaximumHeight(200)
        
    def add_image_item(self, image_path, step_number, step_name=None):
        """Dodaje element z miniaturą i numerem kroku"""
        item = QListWidgetItem()
        
        # Tekst z numerem kroku
        text = f"Krok {step_number}"
        if step_name and step_name != f"Krok {step_number}":
            text = f"Krok {step_number}\n{step_name}"
        
        item.setText(text)
        item.setTextAlignment(Qt.AlignCenter)
        item.setFont(QFont("Arial", 9))
        
        # Utwórz miniaturę
        pixmap = QPixmap(image_path)
        if not pixmap.isNull():
            thumbnail = pixmap.scaled(80, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            item.setIcon(QIcon(thumbnail))
        
        item.setData(Qt.UserRole, image_path)  # Przechowuj ścieżkę do obrazu
        item.setData(Qt.UserRole + 1, step_number)  # Przechowuj numer kroku
        self.addItem(item)
        return item
    
    def update_item(self, index, image_path, step_number, step_name=None):
        """Aktualizuje element listy"""
        if 0 <= index < self.count():
            item = self.item(index)
            
            # Aktualizuj tekst
            text = f"Krok {step_number}"
            if step_name and step_name != f"Krok {step_number}":
                text = f"Krok {step_number}\n{step_name}"
            item.setText(text)
            
            # Aktualizuj miniaturę
            pixmap = QPixmap(image_path)
            if not pixmap.isNull():
                thumbnail = pixmap.scaled(80, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                item.setIcon(QIcon(thumbnail))
            
            item.setData(Qt.UserRole, image_path)
            item.setData(Qt.UserRole + 1, step_number)


class ImageResizerThread(QThread):
    """Wątek do przetwarzania obrazów"""
    progress = Signal(int, int, str)  # current, total, filename
    finished = Signal(list)  # list of processed paths
    
    def __init__(self, image_paths):
        super().__init__()
        self.image_paths = image_paths
        self.temp_dir = tempfile.mkdtemp(prefix="doc_generator_")
        
    def run(self):
        processed_paths = []
        total = len(self.image_paths)
        
        for i, path in enumerate(self.image_paths):
            try:
                processed = self.resize_image(path)
                processed_paths.append(processed)
            except Exception as e:
                processed_paths.append(path)  # Użyj oryginalnego w przypadku błędu
                print(f"Błąd przetwarzania obrazu {path}: {e}")
            
            self.progress.emit(i + 1, total, os.path.basename(path))
        
        self.finished.emit(processed_paths)
    
    def resize_image(self, input_path):
        """Zmniejsza rozmiar obrazu do maksymalnie 800KB"""
        try:
            with Image.open(input_path) as img:
                # Konwertuj do RGB jeśli to konieczne
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                file_size = os.path.getsize(input_path)
                
                if file_size <= 800 * 1024:
                    return input_path
                
                ratio = min(0.8, (800 * 1024) / file_size)
                new_width = int(img.width * ratio)
                new_height = int(img.height * ratio)
                
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                timestamp = int(datetime.now().timestamp() * 1000)
                unique_id = f"{timestamp}_{hashlib.md5(input_path.encode()).hexdigest()[:8]}"
                output_path = os.path.join(self.temp_dir, f"{unique_id}.jpg")
                
                img.save(output_path, 'JPEG', quality=85, optimize=True)
                
                return output_path
                
        except Exception as e:
            print(f"Błąd w resize_image: {e}")
            return input_path
    
    def __del__(self):
        """Sprzątanie tymczasowego katalogu"""
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except:
            pass


class StepPreviewWidget(QWidget):
    """Widget podglądu kroku - tylko obraz"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        
    def setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Etykieta obrazu
        self.image_label = QLabel("Wybierz krok do podglądu")
        self.image_label.setAlignment(Qt.AlignCenter)
        self.image_label.setMinimumHeight(350)
        self.image_label.setFrameStyle(QFrame.Box)
        self.image_label.setStyleSheet("""
            QLabel {
                background-color: white;
                border: 2px solid #cccccc;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
        """)
        
        layout.addWidget(self.image_label)
    
    def set_image(self, image_path):
        """Ustawia obraz w podglądzie"""
        if image_path and os.path.exists(image_path):
            pixmap = QPixmap(image_path)
            if not pixmap.isNull():
                # Skaluj obraz do rozmiaru widgetu
                scaled_pixmap = pixmap.scaled(
                    self.image_label.size() - QSize(40, 40),
                    Qt.KeepAspectRatio,
                    Qt.SmoothTransformation
                )
                self.image_label.setPixmap(scaled_pixmap)
                self.image_label.setText("")
            else:
                self.image_label.setText("Błąd ładowania obrazu")
                self.image_label.setPixmap(QPixmap())
        else:
            self.image_label.setText("Brak obrazu")
            self.image_label.setPixmap(QPixmap())
    
    def resizeEvent(self, event):
        """Obsługa zmiany rozmiaru - przeskalowanie obrazu"""
        super().resizeEvent(event)
        if hasattr(self, 'current_image_path'):
            self.set_image(self.current_image_path)


class GeneratorDokumentow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Piorun 5.1")
        self.setGeometry(100, 100, 1400, 900)
        
        self.ilustracje = []  # Lista ścieżek do obrazów
        self.opisy_krokow = {}  # Słownik opisów kroków {index: {nazwa, opis}}
        self.sciezka_doc = "dokumentacja_kpk.docx"
        self.aktualny_wybrany_krok = None
        
        # Tymczasowy katalog dla przetworzonych obrazów
        self.temp_dir = tempfile.mkdtemp(prefix="doc_generator_")
        
        # Stos do cofania/przywracania zmian
        self.stan_historia = []
        self.aktualny_stan_index = -1
        self.maks_historia = 20
        
        # Motyw
        self.dark_theme = False
        
        # Wielojęzyczność
        self.jezyk = "polski"  # Domyślny język
        self.tlumaczenia = self._inicjalizuj_tlumaczenia()
        
        # Autozapis
        self.autosave_timer = QTimer()
        self.autosave_timer.timeout.connect(self.autozapisz_projekt)
        self.autosave_interval = 300000  # 5 minut w milisekundach
        self.autosave_enabled = True
        self.autosave_dir = os.path.join(os.path.expanduser("~"), "Piorun_autosave")
        os.makedirs(self.autosave_dir, exist_ok=True)
        
        self.setup_ui()
        self.zapisz_stan()  # Zapisz początkowy stan
        
        # Timer do aktualizacji podglądu na żywo
        self.preview_timer = QTimer()
        self.preview_timer.setInterval(500)  # 500 ms
        self.preview_timer.timeout.connect(self.aktualizuj_podglad_opisu_na_zywo)
        
        # Uruchom timer autozapisu
        self.autosave_timer.start(self.autosave_interval)
        
    def __del__(self):
        """Sprzątanie tymczasowych plików przy usuwaniu obiektu"""
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except:
            pass
    
    def _inicjalizuj_tlumaczenia(self):
        """Inicjalizuje słownik tłumaczeń"""
        return {
            "polski": {
                "title": "Piorun 5.1",
                "dane_dokumentu": "Dane dokumentu",
                "kod_produktu": "Kod produktu:",
                "nazwa_dokumentu": "Nazwa dokumentu:",
                "data_dokumentu": "Data dokumentu:",
                "autor_dokumentu": "Autor dokumentu:",
                "zarzadzanie_krokami": "Zarządzanie krokami",
                "lista_krokow": "Lista kroków procedury:",
                "dodaj_ilustracje": "Dodaj ilustracje",
                "wymiana_ilustracji": "Wymiana ilustracji",
                "usun_zaznaczone": "Usuń zaznaczone",
                "edycja_kroku": "Edycja kroku",
                "edytujesz_krok": "Edytujesz krok:",
                "brak_zaznaczenia": "Brak zaznaczenia",
                "nazwa_kroku": "Nazwa kroku:",
                "szczegolowy_opis": "Szczegółowy opis:",
                "zapisz_opis_kroku": "Zapisz opis kroku",
                "podglad": "Podgląd",
                "obraz": "Obraz:",
                "wybierz_krok": "Wybierz krok do podglądu",
                "opcje_dokumentu": "Opcje dokumentu",
                "uklad_krokow": "Układ kroków:",
                "obraz_nad_opisem": "Obraz nad opisem",
                "obraz_pod_opisem": "Obraz pod opisem",
                "obraz_z_lewej": "Obraz z lewej",
                "rozmiar_ilustracji": "Rozmiar ilustracji (cm):",
                "czcionka": "Czcionka:",
                "rozmiar_czcionki": "Rozmiar czcionki:",
                "generuj_otworz": "WYGENERUJ DOKUMENTACJĘ",
                "zapisz_projekt": "Zapisz projekt",
                "wczytaj_projekt": "Wczytaj projekt",
                "wyczyść_wszystko": "Wyczyść wszystko",
                "cofnij": "Cofnij",
                "przywroc": "Przywróć",
                "edytuj_w_paincie": "Edytuj w Paincie",
                "jezyk": "Język:",
                "polski": "Polski",
                "angielski": "Angielski",
                "domyslny_szablon": """1. Po pierwsze:
• Szczegóły pierwszego podpunktu
2. Po drugie:  
• Szczegóły drugiego podpunktu
3. Po trzecie:
• Szczegóły trzeciego podpunktu

Maszyna: """,
                "dokument_tytul": "DOKUMENTACJA KROK PO KROKU",
                "spis_tresci": "Spis kroków procedury",
                "procedura": "Procedura wykonania - krok po kroku",
                "koniec": "Koniec dokumentacji",
                "data_generacji": "Data generacji:",
                "dokumentacja_wygenerowana": "Dokumentacja została wygenerowana automatycznie.",
                "instrukcja_wykonania": "Instrukcja wykonania:",
                "nie_podano": "Nie podano",
                "ostrzeżenie": "Ostrzeżenie",
                "sukces": "Sukces",
                "błąd": "Błąd",
                "dodaj_przynajmniej_jedna_ilustracje": "Dodaj przynajmniej jedną ilustrację!",
                "wprowadz_kod_i_nazwe_dokumentu": "Wprowadź kod i nazwę dokumentu!",
                "błąd_przetwarzania_obrazu": "Błąd przetwarzania obrazu",
                "najpierw_wybierz_ilustracje_do_wymiany": "Najpierw wybierz ilustrację do wymiany!",
                "wybierz_nowa_ilustracje": "Wybierz nową ilustrację",
                "wybierz_ilustracje_krokow_procedury": "Wybierz ilustracje kroków procedury",
                "ilustracja_zostala_wymieniona": "Ilustracja została wymieniona!",
                "najpierw_wybierz_krok": "Najpierw wybierz krok!",
                "opis_kroku_zostal_zapisany": "Opis kroku został zapisany!",
                "błąd_ładowania_obrazu": "Błąd ładowania obrazu",
                "dokumentacja_wygenerowana_i_otwarta": "Dokumentacja została wygenerowana i otwarta",
                "dokument_wygenerowany_ale_nie_otwarty": "Dokument został wygenerowany, ale nie udało się go otworzyć.",
                "wystąpił_błąd_podczas_generowania": "Wystąpił błąd podczas generowania",
                "brak_danych_do_zapisania": "Brak danych do zapisania!",
                "błąd_odczytu_obrazu": "Błąd odczytu obrazu",
                "projekt_zapisany": "Projekt zapisany",
                "obrazy_zostaly_osadzone": "Obrazy zostały osadzone w pliku projektu.",
                "nie_udało_się_zapisać_projektu": "Nie udało się zapisać projektu",
                "projekt_wczytany": "Projekt wczytany",
                "załadowano_obrazy": "Załadowano",
                "obrazów": "obrazów.",
                "nie_udało_się_wczytać_projektu": "Nie udało się wczytać projektu",
                "dokument_nie_istnieje": "Dokument nie istnieje! Wygeneruj go najpierw.",
                "edytowanie_w_paincie": "Edytowanie w Paincie",
                "czy_zapisac_zmiany": "Czy zapisać zmiany w obrazie? Zamknij Paint, a następnie kliknij 'Tak' aby zaktualizować obraz w aplikacji.",
                "funkcja_dostepna_tylko_windows": "Funkcja edycji w Paincie dostępna tylko w systemie Windows.",
                "błąd_otwierania_paint": "Błąd otwierania Paint",
                "obraz_zaktualizowany": "Obraz został zaktualizowany!",
                "najpierw_wybierz_obraz": "Najpierw wybierz obraz do edycji!",
                "generowanie_dokumentu": "Generowanie dokumentu...",
                "krok": "Krok",
                "liczba_krokow": "Liczba kroków:",
                "wybierz_katalog": "Wybierz katalog do zapisu dokumentu",
                "utworz_kopie_zapasowa": "Utwórz kopię zapasową",
                "kopia_zapasowa_utworzona": "Kopia zapasowa utworzona",
                "katalog_dla_kopii_zapasowej": "Wybierz katalog dla kopii zapasowej",
                "motyw": "Motyw",
                "jasny_motyw": "Jasny motyw",
                "ciemny_motyw": "Ciemny motyw"
            },
            "angielski": {
                "title": "Piorun 5.1",
                "dane_dokumentu": "Document Data",
                "kod_produktu": "Product code:",
                "nazwa_dokumentu": "Document name:",
                "data_dokumentu": "Document date:",
                "autor_dokumentu": "Document author:",
                "zarzadzanie_krokami": "Step Management",
                "lista_krokow": "Procedure steps list:",
                "dodaj_ilustracje": "Add illustrations",
                "wymiana_ilustracji": "Replace illustration",
                "usun_zaznaczone": "Delete selected",
                "edycja_kroku": "Step editing",
                "edytujesz_krok": "You are editing step:",
                "brak_zaznaczenia": "No selection",
                "nazwa_kroku": "Step name:",
                "szczegolowy_opis": "Detailed description:",
                "zapisz_opis_kroku": "Save step description",
                "podglad": "Preview",
                "obraz": "Image:",
                "wybierz_krok": "Select step for preview",
                "opcje_dokumentu": "Document options",
                "uklad_krokow": "Step layout:",
                "obraz_nad_opisem": "Image above description",
                "obraz_pod_opisem": "Image below description",
                "obraz_z_lewej": "Image on the left",
                "rozmiar_ilustracji": "Illustration size (cm):",
                "czcionka": "Font:",
                "rozmiar_czcionki": "Font size:",
                "generuj_otworz": "GENERATE DOCUMENTATION",
                "zapisz_projekt": "Save project",
                "wczytaj_projekt": "Load project",
                "wyczyść_wszystko": "Clear all",
                "cofnij": "Undo",
                "przywroc": "Redo",
                "edytuj_w_paincie": "Edit in Paint",
                "jezyk": "Language:",
                "polski": "Polish",
                "angielski": "English",
                "domyslny_szablon": """1. First:
• Details of the first point
2. Second:  
• Details of the second point
3. Third:
• Details of the third point

Machine: [Enter machine or device name here]""",
                "dokument_tytul": "STEP BY STEP DOCUMENTATION",
                "spis_tresci": "Procedure steps table of contents",
                "procedura": "Step by step procedure",
                "koniec": "End of documentation",
                "data_generacji": "Generation date:",
                "dokumentacja_wygenerowana": "Documentation has been generated automatically.",
                "instrukcja_wykonania": "Instruction:",
                "nie_podano": "Not provided",
                "ostrzeżenie": "Warning",
                "sukces": "Success",
                "błąd": "Error",
                "dodaj_przynajmniej_jedna_ilustracje": "Add at least one illustration!",
                "wprowadz_kod_i_nazwe_dokumentu": "Enter document code and name!",
                "błąd_przetwarzania_obrazu": "Image processing error",
                "najpierw_wybierz_ilustracje_do_wymiany": "First select an illustration to replace!",
                "wybierz_nowa_ilustracje": "Select new illustration",
                "wybierz_ilustracje_krokow_procedury": "Select procedure step illustrations",
                "ilustracja_zostala_wymieniona": "Illustration has been replaced!",
                "najpierw_wybierz_krok": "First select a step!",
                "opis_kroku_zostal_zapisany": "Step description has been saved!",
                "błąd_ładowania_obrazu": "Error loading image",
                "dokumentacja_wygenerowana_i_otwarta": "Documentation has been generated and opened",
                "dokument_wygenerowany_ale_nie_otwarty": "Document has been generated but could not be opened.",
                "wystąpił_błąd_podczas_generowania": "An error occurred during generation",
                "brak_danych_do_zapisania": "No data to save!",
                "błąd_odczytu_obrazu": "Error reading image",
                "projekt_zapisany": "Project saved",
                "obrazy_zostaly_osadzone": "Images have been embedded in the project file.",
                "nie_udało_się_zapisać_projektu": "Failed to save project",
                "projekt_wczytany": "Project loaded",
                "załadowano_obrazy": "Loaded",
                "obrazów": "images.",
                "nie_udało_się_wczytać_projektu": "Failed to load project",
                "dokument_nie_istnieje": "Document does not exist! Generate it first.",
                "edytowanie_w_paincie": "Editing in Paint",
                "czy_zapisac_zmiany": "Save changes to the image? Close Paint, then click 'Yes' to update the image in the application.",
                "funkcja_dostepna_tylko_windows": "Paint editing function available only on Windows.",
                "błąd_otwierania_paint": "Error opening Paint",
                "obraz_zaktualizowany": "Image has been updated!",
                "najpierw_wybierz_obraz": "First select an image to edit!",
                "generowanie_dokumentu": "Generating document...",
                "krok": "Step",
                "liczba_krokow": "Number of steps:",
                "wybierz_katalog": "Select directory to save document",
                "utworz_kopie_zapasowa": "Create backup",
                "kopia_zapasowa_utworzona": "Backup created",
                "katalog_dla_kopii_zapasowej": "Select directory for backup",
                "motyw": "Theme",
                "jasny_motyw": "Light theme",
                "ciemny_motyw": "Dark theme"
            }
        }
    
    def t(self, klucz):
        """Zwraca tłumaczenie dla danego klucza w aktualnym języku"""
        tlumaczenia = self.tlumaczenia.get(self.jezyk, {})
        return tlumaczenia.get(klucz, f"[{klucz}]")
    
    def setup_ui(self):
        """Tworzy interfejs użytkownika"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(5, 5, 5, 5)
        main_layout.setSpacing(5)
        
        # Górna część - dane dokumentu i zarządzanie krokami
        top_splitter = QSplitter(Qt.Horizontal)
        
        # Lewa strona - dane dokumentu
        data_group = QGroupBox(self.t("dane_dokumentu"))
        data_layout = QGridLayout()
        
        data_layout.addWidget(QLabel(self.t("kod_produktu")), 0, 0)
        self.kod_edit = QLineEdit("xxx-xxxx-xxx")
        data_layout.addWidget(self.kod_edit, 0, 1)
        
        data_layout.addWidget(QLabel(self.t("nazwa_dokumentu")), 1, 0)
        self.nazwa_edit = QLineEdit("Dokumentacja produktu")
        data_layout.addWidget(self.nazwa_edit, 1, 1)
        
        data_layout.addWidget(QLabel(self.t("data_dokumentu")), 2, 0)
        self.data_edit = QLineEdit(datetime.now().strftime("%Y-%m-%d"))
        data_layout.addWidget(self.data_edit, 2, 1)
        
        data_layout.addWidget(QLabel(self.t("autor_dokumentu")), 3, 0)
        self.autor_edit = QLineEdit("Dorota Zaręba")  # DOMYŚLNIE
        data_layout.addWidget(self.autor_edit, 3, 1)
        
        data_group.setLayout(data_layout)
        top_splitter.addWidget(data_group)
        
        # Prawa strona - zarządzanie krokami
        steps_group = QGroupBox(self.t("zarzadzanie_krokami"))
        steps_layout = QVBoxLayout()
        
        steps_layout.addWidget(QLabel(self.t("lista_krokow")))
        
        # Przyciski zarządzania
        btn_layout = QHBoxLayout()
        
        self.add_btn = QPushButton(self.t("dodaj_ilustracje"))
        self.add_btn.clicked.connect(self.dodaj_ilustracje)
        btn_layout.addWidget(self.add_btn)
        
        self.replace_btn = QPushButton(self.t("wymiana_ilustracji"))
        self.replace_btn.clicked.connect(self.wymien_ilustracje)
        btn_layout.addWidget(self.replace_btn)
        
        self.delete_btn = QPushButton(self.t("usun_zaznaczone"))
        self.delete_btn.clicked.connect(self.usun_ilustracje)
        btn_layout.addWidget(self.delete_btn)
        
        # Przyciski przenoszenia
        move_layout = QHBoxLayout()
        self.move_up_btn = QPushButton("▲")
        self.move_up_btn.clicked.connect(self.przenies_w_gore)
        self.move_up_btn.setFixedWidth(30)
        move_layout.addWidget(self.move_up_btn)
        
        self.move_down_btn = QPushButton("▼")
        self.move_down_btn.clicked.connect(self.przenies_w_dol)
        self.move_down_btn.setFixedWidth(30)
        move_layout.addWidget(self.move_down_btn)
        
        move_layout.addStretch()
        btn_layout.addLayout(move_layout)
        
        steps_layout.addLayout(btn_layout)
        
        # Lista kroków Z MINIATURAMI
        self.steps_list = ThumbnailListWidget()
        self.steps_list.itemClicked.connect(self.pokaz_podglad)
        steps_layout.addWidget(self.steps_list)
        
        steps_group.setLayout(steps_layout)
        top_splitter.addWidget(steps_group)
        
        top_splitter.setSizes([400, 600])
        main_layout.addWidget(top_splitter)
        
        # Środkowa część - edycja i podgląd
        middle_splitter = QSplitter(Qt.Horizontal)
        
        # Lewa strona - edycja kroku
        edit_group = QGroupBox(self.t("edycja_kroku"))
        edit_layout = QVBoxLayout()
        
        # Info o kroku
        self.step_info_label = QLabel(self.t("brak_zaznaczenia"))
        self.step_info_label.setStyleSheet("color: red; font-weight: bold; font-size: 12px;")
        edit_layout.addWidget(self.step_info_label)
        
        edit_layout.addWidget(QLabel(self.t("nazwa_kroku")))
        self.step_name_edit = QLineEdit()
        self.step_name_edit.textChanged.connect(self.aktualizuj_nazwe_kroku_na_liscie)
        edit_layout.addWidget(self.step_name_edit)
        
        edit_layout.addWidget(QLabel(self.t("szczegolowy_opis")))
        self.step_desc_edit = QTextEdit()
        self.step_desc_edit.setPlaceholderText(self.t("domyslny_szablon"))
        self.step_desc_edit.textChanged.connect(self.start_preview_timer)
        edit_layout.addWidget(self.step_desc_edit)
        
        self.save_step_btn = QPushButton(self.t("zapisz_opis_kroku"))
        self.save_step_btn.clicked.connect(self.zapisz_opis_kroku)
        edit_layout.addWidget(self.save_step_btn)
        
        edit_group.setLayout(edit_layout)
        middle_splitter.addWidget(edit_group)
        
        # Prawa strona - podgląd (TYLKO OBRAZ)
        preview_group = QGroupBox(self.t("podglad"))
        preview_layout = QVBoxLayout()
        
        self.preview_widget = StepPreviewWidget()
        preview_layout.addWidget(self.preview_widget)
        
        preview_group.setLayout(preview_layout)
        middle_splitter.addWidget(preview_group)
        
        middle_splitter.setSizes([400, 600])
        main_layout.addWidget(middle_splitter)
        
        # Dolna część - opcje dokumentu
        options_group = QGroupBox(self.t("opcje_dokumentu"))
        options_layout = QHBoxLayout()
        
        # Układ kroków
        layout_box = QVBoxLayout()
        layout_box.addWidget(QLabel(self.t("uklad_krokow")))
        self.layout_combo = QComboBox()
        self.layout_combo.addItem(self.t("obraz_z_lewej"), "lewo_prawo")
        self.layout_combo.addItem(self.t("obraz_nad_opisem"), "gora")
        self.layout_combo.addItem(self.t("obraz_pod_opisem"), "dol")
        layout_box.addWidget(self.layout_combo)
        options_layout.addLayout(layout_box)
        
        # Rozmiar ilustracji
        size_box = QVBoxLayout()
        size_box.addWidget(QLabel(self.t("rozmiar_ilustracji")))
        self.size_combo = QComboBox()
        self.size_combo.addItems(["2", "3", "5", "6", "8", "10", "12"])
        self.size_combo.setCurrentText("8")
        size_box.addWidget(self.size_combo)
        options_layout.addLayout(size_box)
        
        # Czcionka
        font_box = QVBoxLayout()
        font_box.addWidget(QLabel(self.t("czcionka")))
        self.font_combo = QComboBox()
        self.font_combo.addItems(["Arial", "Times New Roman", "Calibri", "Verdana"])
        self.font_combo.setCurrentText("Arial")
        font_box.addWidget(self.font_combo)
        options_layout.addLayout(font_box)
        
        # Rozmiar czcionki
        font_size_box = QVBoxLayout()
        font_size_box.addWidget(QLabel(self.t("rozmiar_czcionki")))
        self.font_size_combo = QComboBox()
        self.font_size_combo.addItems(["9", "10", "11", "12", "14", "16"])
        self.font_size_combo.setCurrentText("11")
        font_size_box.addWidget(self.font_size_combo)
        options_layout.addLayout(font_size_box)
        
        # Język
        lang_box = QVBoxLayout()
        lang_box.addWidget(QLabel(self.t("jezyk")))
        self.lang_combo = QComboBox()
        self.lang_combo.addItem(self.t("polski"), "polski")
        self.lang_combo.addItem(self.t("angielski"), "angielski")
        self.lang_combo.currentIndexChanged.connect(self.zmien_jezyk)
        lang_box.addWidget(self.lang_combo)
        options_layout.addLayout(lang_box)
        
        options_layout.addStretch()
        options_group.setLayout(options_layout)
        main_layout.addWidget(options_group)
        
        # PRZYCISKI AKCJI - WSZYSTKIE W JEDNEJ LINII
        action_layout = QHBoxLayout()
        
        # Duży przycisk generowania dokumentacji (20% mniejszy)
        self.generate_btn = QPushButton(self.t("generuj_otworz"))
        self.generate_btn.clicked.connect(self.generuj_i_otworz_dokument)
        self.generate_btn.setFixedWidth(240)  # 20% mniejszy niż 300px
        self.generate_btn.setStyleSheet("""
            QPushButton {
                font-size: 14px;
                font-weight: bold;
                padding: 10px 20px;
                background-color: #4CAF50;
                color: white;
                border: 2px solid #388E3C;
                border-radius: 6px;
                min-height: 40px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:pressed {
                background-color: #388E3C;
            }
        """)
        action_layout.addWidget(self.generate_btn)
        
        # Pozostałe przyciski
        action_layout.addWidget(QLabel("   "))
        
        self.save_project_btn = QPushButton(self.t("zapisz_projekt"))
        self.save_project_btn.clicked.connect(self.zapisz_projekt)
        action_layout.addWidget(self.save_project_btn)
        
        self.load_project_btn = QPushButton(self.t("wczytaj_projekt"))
        self.load_project_btn.clicked.connect(self.wczytaj_projekt)
        action_layout.addWidget(self.load_project_btn)
        
        self.clear_btn = QPushButton(self.t("wyczyść_wszystko"))
        self.clear_btn.clicked.connect(self.wyczysc_wszystko)
        action_layout.addWidget(self.clear_btn)
        
        self.backup_btn = QPushButton(self.t("utworz_kopie_zapasowa"))
        self.backup_btn.clicked.connect(self.utworz_kopie_zapasowa)
        action_layout.addWidget(self.backup_btn)
        
        action_layout.addWidget(QLabel("   "))
        
        self.undo_btn = QPushButton(self.t("cofnij"))
        self.undo_btn.clicked.connect(self.cofnij)
        action_layout.addWidget(self.undo_btn)
        
        self.redo_btn = QPushButton(self.t("przywroc"))
        self.redo_btn.clicked.connect(self.przywroc)
        action_layout.addWidget(self.redo_btn)
        
        self.paint_btn = QPushButton(self.t("edytuj_w_paincie"))
        self.paint_btn.clicked.connect(self.edytuj_w_paincie)
        action_layout.addWidget(self.paint_btn)
        
        action_layout.addStretch()
        
        main_layout.addLayout(action_layout)
        
        # Menu
        self.setup_menu()
        
        # Ustaw domyślny motyw (jasny)
        self.zmien_motyw(False)
    
    def setup_menu(self):
        """Tworzy menu aplikacji"""
        menubar = self.menuBar()
        
        # Menu plik
        file_menu = menubar.addMenu("Plik")
        
        new_action = QAction("Nowy projekt", self)
        new_action.triggered.connect(self.wyczysc_wszystko)
        file_menu.addAction(new_action)
        
        load_action = QAction("Wczytaj projekt", self)
        load_action.triggered.connect(self.wczytaj_projekt)
        file_menu.addAction(load_action)
        
        save_action = QAction("Zapisz projekt", self)
        save_action.triggered.connect(self.zapisz_projekt)
        file_menu.addAction(save_action)
        
        file_menu.addSeparator()
        
        generate_action = QAction("Generuj dokumentację", self)
        generate_action.triggered.connect(self.generuj_i_otworz_dokument)
        file_menu.addAction(generate_action)
        
        file_menu.addSeparator()
        
        # Autozapis
        autosave_menu = file_menu.addMenu("Autozapis")
        
        enable_autosave_action = QAction("Włącz autozapis", self)
        enable_autosave_action.setCheckable(True)
        enable_autosave_action.setChecked(True)
        enable_autosave_action.triggered.connect(
            lambda checked: setattr(self, 'autosave_enabled', checked)
        )
        autosave_menu.addAction(enable_autosave_action)
        
        autosave_menu.addSeparator()
        
        open_autosave_dir_action = QAction("Otwórz folder autozapisu", self)
        open_autosave_dir_action.triggered.connect(self.otworz_folder_autozapisu)
        autosave_menu.addAction(open_autosave_dir_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction("Zakończ", self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Menu widok
        view_menu = menubar.addMenu("Widok")
        
        # Motywy
        theme_menu = view_menu.addMenu(self.t("motyw"))
        
        light_theme_action = QAction(self.t("jasny_motyw"), self)
        light_theme_action.triggered.connect(lambda: self.zmien_motyw(False))
        theme_menu.addAction(light_theme_action)
        
        dark_theme_action = QAction(self.t("ciemny_motyw"), self)
        dark_theme_action.triggered.connect(lambda: self.zmien_motyw(True))
        theme_menu.addAction(dark_theme_action)
        
        # Menu język
        lang_menu = menubar.addMenu(self.t("jezyk"))
        
        polish_action = QAction("Polski", self)
        polish_action.triggered.connect(lambda: self.zmien_jezyk_combo("polski"))
        lang_menu.addAction(polish_action)
        
        english_action = QAction("English", self)
        english_action.triggered.connect(lambda: self.zmien_jezyk_combo("angielski"))
        lang_menu.addAction(english_action)
    
    def otworz_folder_autozapisu(self):
        """Otwiera folder z autozapisami w eksploratorze plików"""
        try:
            if sys.platform == "win32":
                os.startfile(self.autosave_dir)
            elif sys.platform == "darwin":
                subprocess.call(["open", self.autosave_dir])
            else:
                subprocess.call(["xdg-open", self.autosave_dir])
        except Exception as e:
            QMessageBox.warning(self, "Błąd", f"Nie można otworzyć folderu: {e}")
    
    def zmien_jezyk_combo(self, jezyk):
        """Zmienia język z menu"""
        self.jezyk = jezyk
        index = self.lang_combo.findData(jezyk)
        if index >= 0:
            self.lang_combo.setCurrentIndex(index)
        self.odswiez_interfejs()
    
    def zmien_motyw(self, dark=False):
        """Zmienia motyw aplikacji"""
        self.dark_theme = dark
        
        if dark:
            # Ciemny motyw
            palette = QPalette()
            palette.setColor(QPalette.Window, QColor(53, 53, 53))
            palette.setColor(QPalette.WindowText, Qt.white)
            palette.setColor(QPalette.Base, QColor(25, 25, 25))
            palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
            palette.setColor(QPalette.ToolTipBase, Qt.white)
            palette.setColor(QPalette.ToolTipText, Qt.white)
            palette.setColor(QPalette.Text, Qt.white)
            palette.setColor(QPalette.Button, QColor(53, 53, 53))
            palette.setColor(QPalette.ButtonText, Qt.white)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Link, QColor(42, 130, 218))
            palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
            palette.setColor(QPalette.HighlightedText, Qt.black)
            
            # Styl dla listy
            self.steps_list.setStyleSheet("""
                QListWidget {
                    background-color: #2d2d2d;
                    border: 1px solid #444444;
                    border-radius: 4px;
                }
                QListWidget::item {
                    background-color: #3d3d3d;
                    border: 1px solid #555555;
                    border-radius: 4px;
                    padding: 5px;
                    color: white;
                }
                QListWidget::item:selected {
                    background-color: #4CAF50;
                    border: 2px solid #388E3C;
                    color: white;
                }
            """)
            
            # Styl dla podglądu
            self.preview_widget.image_label.setStyleSheet("""
                QLabel {
                    background-color: #2d2d2d;
                    border: 2px solid #444444;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                    color: white;
                }
            """)
            
            # Styl dla dużego przycisku
            self.generate_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px 20px;
                    background-color: #4CAF50;
                    color: white;
                    border: 2px solid #388E3C;
                    border-radius: 6px;
                    min-height: 40px;
                }
                QPushButton:hover {
                    background-color: #45a049;
                }
                QPushButton:pressed {
                    background-color: #388E3C;
                }
            """)
            
        else:
            # Jasny motyw
            palette = QPalette()
            palette.setColor(QPalette.Window, QColor(240, 240, 240))
            palette.setColor(QPalette.WindowText, Qt.black)
            palette.setColor(QPalette.Base, QColor(255, 255, 255))
            palette.setColor(QPalette.AlternateBase, QColor(240, 240, 240))
            palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
            palette.setColor(QPalette.ToolTipText, Qt.black)
            palette.setColor(QPalette.Text, Qt.black)
            palette.setColor(QPalette.Button, QColor(240, 240, 240))
            palette.setColor(QPalette.ButtonText, Qt.black)
            palette.setColor(QPalette.BrightText, Qt.red)
            palette.setColor(QPalette.Link, QColor(42, 130, 218))
            palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
            palette.setColor(QPalette.HighlightedText, Qt.white)
            
            # Styl dla listy
            self.steps_list.setStyleSheet("""
                QListWidget {
                    background-color: #f0f0f0;
                    border: 1px solid #cccccc;
                    border-radius: 4px;
                }
                QListWidget::item {
                    background-color: white;
                    border: 1px solid #dddddd;
                    border-radius: 4px;
                    padding: 5px;
                    color: black;
                }
                QListWidget::item:selected {
                    background-color: #4CAF50;
                    border: 2px solid #388E3C;
                    color: white;
                }
            """)
            
            # Styl dla podglądu
            self.preview_widget.image_label.setStyleSheet("""
                QLabel {
                    background-color: white;
                    border: 2px solid #cccccc;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                    color: black;
                }
            """)
            
            # Styl dla dużego przycisku
            self.generate_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px;
                    font-weight: bold;
                    padding: 10px 20px;
                    background-color: #4CAF50;
                    color: white;
                    border: 2px solid #388E3C;
                    border-radius: 6px;
                    min-height: 40px;
                }
                QPushButton:hover {
                    background-color: #45a049;
                }
                QPushButton:pressed {
                    background-color: #388E3C;
                }
            """)
        
        QApplication.instance().setPalette(palette)
    
    def start_preview_timer(self):
        """Uruchamia timer do aktualizacji podglądu na żywo"""
        self.preview_timer.start()
    
    def zapisz_stan(self):
        """Zapisuje aktualny stan do historii"""
        stan = {
            'ilustracje': self.ilustracje.copy(),
            'opisy_krokow': copy.deepcopy(self.opisy_krokow),
            'aktualny_wybrany_krok': self.aktualny_wybrany_krok,
            'kod': self.kod_edit.text(),
            'nazwa': self.nazwa_edit.text(),
            'data': self.data_edit.text(),
            'autor': self.autor_edit.text()
        }
        
        if len(self.stan_historia) >= self.maks_historia:
            self.stan_historia.pop(0)
            if self.aktualny_stan_index > 0:
                self.aktualny_stan_index -= 1
        
        self.stan_historia.append(stan)
        self.aktualny_stan_index = len(self.stan_historia) - 1
    
    def cofnij(self):
        """Cofa ostatnią zmianę"""
        if self.aktualny_stan_index > 0:
            self.aktualny_stan_index -= 1
            stan = self.stan_historia[self.aktualny_stan_index]
            
            self.ilustracje = stan['ilustracje'].copy()
            self.opisy_krokow = copy.deepcopy(stan['opisy_krokow'])
            self.aktualny_wybrany_krok = stan['aktualny_wybrany_krok']
            self.kod_edit.setText(stan['kod'])
            self.nazwa_edit.setText(stan['nazwa'])
            self.data_edit.setText(stan['data'])
            self.autor_edit.setText(stan['autor'])
            
            self.odswiez_liste()
            if self.aktualny_wybrany_krok is not None:
                self.steps_list.setCurrentRow(self.aktualny_wybrany_krok)
                self.pokaz_podglad()
            self.aktualizuj_info_o_kroku()
    
    def przywroc(self):
        """Przywraca ostatnio cofniętą zmianę"""
        if self.aktualny_stan_index < len(self.stan_historia) - 1:
            self.aktualny_stan_index += 1
            stan = self.stan_historia[self.aktualny_stan_index]
            
            self.ilustracje = stan['ilustracje'].copy()
            self.opisy_krokow = copy.deepcopy(stan['opisy_krokow'])
            self.aktualny_wybrany_krok = stan['aktualny_wybrany_krok']
            self.kod_edit.setText(stan['kod'])
            self.nazwa_edit.setText(stan['nazwa'])
            self.data_edit.setText(stan['data'])
            self.autor_edit.setText(stan['autor'])
            
            self.odswiez_liste()
            if self.aktualny_wybrany_krok is not None:
                self.steps_list.setCurrentRow(self.aktualny_wybrany_krok)
                self.pokaz_podglad()
            self.aktualizuj_info_o_kroku()
    
    def zmien_jezyk(self):
        """Zmienia język interfejsu"""
        self.jezyk = self.lang_combo.currentData()
        self.odswiez_interfejs()
    
    def odswiez_interfejs(self):
        """Odświeża wszystkie elementy interfejsu w aktualnym języku"""
        self.setWindowTitle(self.t("title"))
        
        # Aktualizuj tytuły grup
        for widget in self.findChildren(QGroupBox):
            title = widget.title()
            if title in ["Dane dokumentu", "Document Data"]:
                widget.setTitle(self.t("dane_dokumentu"))
            elif title in ["Zarządzanie krokami", "Step Management"]:
                widget.setTitle(self.t("zarzadzanie_krokami"))
            elif title in ["Edycja kroku", "Step editing"]:
                widget.setTitle(self.t("edycja_kroku"))
            elif title in ["Podgląd", "Preview"]:
                widget.setTitle(self.t("podglad"))
            elif title in ["Opcje dokumentu", "Document options"]:
                widget.setTitle(self.t("opcje_dokumentu"))
        
        # Aktualizuj etykiety
        for widget in self.findChildren(QLabel):
            text = widget.text()
            if text in ["Kod produktu:", "Product code:"]:
                widget.setText(self.t("kod_produktu"))
            elif text in ["Nazwa dokumentu:", "Document name:"]:
                widget.setText(self.t("nazwa_dokumentu"))
            elif text in ["Data dokumentu:", "Document date:"]:
                widget.setText(self.t("data_dokumentu"))
            elif text in ["Autor dokumentu:", "Document author:"]:
                widget.setText(self.t("autor_dokumentu"))
            elif text in ["Lista kroków procedury:", "Procedure steps list:"]:
                widget.setText(self.t("lista_krokow"))
            elif text in ["Edytujesz krok:", "You are editing step:"]:
                widget.setText(self.t("edytujesz_krok"))
            elif text in ["Nazwa kroku:", "Step name:"]:
                widget.setText(self.t("nazwa_kroku"))
            elif text in ["Szczegółowy opis:", "Detailed description:"]:
                widget.setText(self.t("szczegolowy_opis"))
            elif text in ["Obraz:", "Image:"]:
                widget.setText(self.t("obraz"))
            elif text in ["Układ kroków:", "Step layout:"]:
                widget.setText(self.t("uklad_krokow"))
            elif text in ["Rozmiar ilustracji (cm):", "Illustration size (cm):"]:
                widget.setText(self.t("rozmiar_ilustracji"))
            elif text in ["Czcionka:", "Font:"]:
                widget.setText(self.t("czcionka"))
            elif text in ["Rozmiar czcionki:", "Font size:"]:
                widget.setText(self.t("rozmiar_czcionki"))
            elif text in ["Język:", "Language:"]:
                widget.setText(self.t("jezyk"))
        
        # Aktualizuj przyciski
        for widget in self.findChildren(QPushButton):
            text = widget.text()
            if text in ["Dodaj ilustracje", "Add illustrations"]:
                widget.setText(self.t("dodaj_ilustracje"))
            elif text in ["Wymiana ilustracji", "Replace illustration"]:
                widget.setText(self.t("wymiana_ilustracji"))
            elif text in ["Usuń zaznaczone", "Delete selected"]:
                widget.setText(self.t("usun_zaznaczone"))
            elif text in ["Zapisz opis kroku", "Save step description"]:
                widget.setText(self.t("zapisz_opis_kroku"))
            elif text in ["WYGENERUJ DOKUMENTACJĘ", "GENERATE DOCUMENTATION"]:
                widget.setText(self.t("generuj_otworz"))
            elif text in ["Zapisz projekt", "Save project"]:
                widget.setText(self.t("zapisz_projekt"))
            elif text in ["Wczytaj projekt", "Load project"]:
                widget.setText(self.t("wczytaj_projekt"))
            elif text in ["Wyczyść wszystko", "Clear all"]:
                widget.setText(self.t("wyczyść_wszystko"))
            elif text in ["Cofnij", "Undo"]:
                widget.setText(self.t("cofnij"))
            elif text in ["Przywróć", "Redo"]:
                widget.setText(self.t("przywroc"))
            elif text in ["Edytuj w Paincie", "Edit in Paint"]:
                widget.setText(self.t("edytuj_w_paincie"))
            elif text in ["Utwórz kopię zapasową", "Create backup"]:
                widget.setText(self.t("utworz_kopie_zapasowa"))
        
        # Aktualizuj comboboxy
        self.layout_combo.setItemText(0, self.t("obraz_z_lewej"))
        self.layout_combo.setItemText(1, self.t("obraz_nad_opisem"))
        self.layout_combo.setItemText(2, self.t("obraz_pod_opisem"))
        
        self.lang_combo.setItemText(0, self.t("polski"))
        self.lang_combo.setItemText(1, self.t("angielski"))
        
        # Aktualizuj informację o kroku
        self.aktualizuj_info_o_kroku()
        
        # Aktualizuj placeholder
        self.step_desc_edit.setPlaceholderText(self.t("domyslny_szablon"))
    
    def dodaj_ilustracje(self):
        """Dodaje nowe ilustracje do projektu"""
        files, _ = QFileDialog.getOpenFileNames(
            self,
            self.t("wybierz_ilustracje_krokow_procedury"),
            "",
            "Pliki obrazów (*.jpg *.jpeg *.png *.gif *.bmp *.tiff *.webp);;Wszystkie pliki (*.*)"
        )
        
        if not files:
            return
        
        # Pokaż dialog postępu
        progress = QProgressDialog(
            self.t("generowanie_dokumentu"),
            "Anuluj",
            0,
            len(files),
            self
        )
        progress.setWindowTitle("Przetwarzanie obrazów")
        progress.setWindowModality(Qt.WindowModal)
        progress.show()
        
        # Zapisz stan przed dodaniem
        self.zapisz_stan()
        
        # Użyj wątku do przetwarzania obrazów
        self.resizer_thread = ImageResizerThread(files)
        self.resizer_thread.progress.connect(
            lambda current, total, filename: progress.setValue(current)
        )
        self.resizer_thread.finished.connect(
            lambda processed_paths: self.on_images_processed(processed_paths, progress)
        )
        self.resizer_thread.start()
    
    def on_images_processed(self, processed_paths, progress):
        """Obsługuje zakończenie przetwarzania obrazów"""
        progress.close()
        
        for i, plik in enumerate(processed_paths):
            self.ilustracje.append(plik)
            nazwa_pliku = os.path.basename(plik)
            nazwa_kroku = os.path.splitext(nazwa_pliku)[0]
            
            index = len(self.ilustracje) - 1
            self.opisy_krokow[index] = {
                'nazwa': nazwa_kroku,
                'opis': self.t("domyslny_szablon")
            }
            
            # Dodaj do listy z miniaturą i numerem kroku
            step_number = index + 1
            self.steps_list.add_image_item(plik, step_number, nazwa_kroku)
        
        if processed_paths:
            self.steps_list.setCurrentRow(len(self.ilustracje) - 1)
            self.pokaz_podglad()
    
    def wymien_ilustracje(self):
        """Wymienia zaznaczoną ilustrację na nową"""
        current_row = self.steps_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("najpierw_wybierz_ilustracje_do_wymiany"))
            return
        
        plik, _ = QFileDialog.getOpenFileName(
            self,
            self.t("wybierz_nowa_ilustracje"),
            "",
            "Pliki obrazów (*.jpg *.jpeg *.png *.gif *.bmp *.tiff *.webp);;Wszystkie pliki (*.*)"
        )
        
        if plik:
            self.zapisz_stan()
            index = current_row
            
            # Przetwórz obraz
            przetworzony_plik = self.resize_image(plik)
            self.ilustracje[index] = przetworzony_plik
            
            # Zachowaj starą nazwę jeśli była edytowana
            if index in self.opisy_krokow and self.opisy_krokow[index]['nazwa']:
                zachowana_nazwa = self.opisy_krokow[index]['nazwa']
            else:
                nazwa_pliku = os.path.basename(plik)
                zachowana_nazwa = os.path.splitext(nazwa_pliku)[0]
            
            # Aktualizuj listę z numerem kroku
            step_number = index + 1
            self.steps_list.update_item(index, przetworzony_plik, step_number, zachowana_nazwa)
            
            # Aktualizuj opis
            if index not in self.opisy_krokow:
                self.opisy_krokow[index] = {
                    'nazwa': zachowana_nazwa,
                    'opis': self.t("domyslny_szablon")
                }
            else:
                self.opisy_krokow[index]['nazwa'] = zachowana_nazwa
            
            self.steps_list.setCurrentRow(index)
            self.pokaz_podglad()
            QMessageBox.information(self, self.t("sukces"), 
                                  self.t("ilustracja_zostala_wymieniona"))
    
    def resize_image(self, input_path):
        """Zmniejsza rozmiar obrazu (funkcja pomocnicza)"""
        try:
            with Image.open(input_path) as img:
                # Konwertuj do RGB jeśli to konieczne
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                file_size = os.path.getsize(input_path)
                
                if file_size <= 800 * 1024:
                    return input_path
                
                ratio = min(0.8, (800 * 1024) / file_size)
                new_width = int(img.width * ratio)
                new_height = int(img.height * ratio)
                
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                timestamp = int(datetime.now().timestamp() * 1000)
                unique_id = f"{timestamp}_{hashlib.md5(input_path.encode()).hexdigest()[:8]}"
                output_path = os.path.join(self.temp_dir, f"{unique_id}.jpg")
                
                img.save(output_path, 'JPEG', quality=85, optimize=True)
                
                return output_path
                
        except Exception as e:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              f"{self.t('błąd_przetwarzania_obrazu')} {os.path.basename(input_path)}: {e}")
            return input_path
    
    def usun_ilustracje(self):
        """Usuwa zaznaczoną ilustrację"""
        current_row = self.steps_list.currentRow()
        if current_row < 0:
            return
        
        self.zapisz_stan()
        index = current_row
        
        # Usuń z listy
        item = self.steps_list.takeItem(index)
        if item:
            del item
        
        # Usuń z listy obrazów
        if index < len(self.ilustracje):
            self.ilustracje.pop(index)
        
        # Usuń opis
        if index in self.opisy_krokow:
            del self.opisy_krokow[index]
        
        # Przesuń opisy dla pozostałych kroków
        nowe_opisy = {}
        for i in range(len(self.ilustracje)):
            if i >= index:
                if i+1 in self.opisy_krokow:
                    nowe_opisy[i] = self.opisy_krokow[i+1]
            else:
                if i in self.opisy_krokow:
                    nowe_opisy[i] = self.opisy_krokow[i]
        
        self.opisy_krokow = nowe_opisy
        
        # Odśwież listę z nowymi numerami kroków
        self.odswiez_liste()
        
        # Odśwież podgląd
        if self.steps_list.count() > 0:
            if index >= self.steps_list.count():
                index = self.steps_list.count() - 1
            self.steps_list.setCurrentRow(index)
            self.pokaz_podglad()
        else:
            self.preview_widget.set_image(None)
            self.step_name_edit.clear()
            self.step_desc_edit.clear()
            self.aktualny_wybrany_krok = None
            self.aktualizuj_info_o_kroku()
    
    def przenies_w_gore(self):
        """Przenosi zaznaczony krok w górę"""
        current_row = self.steps_list.currentRow()
        if current_row <= 0:
            return
        
        self.zapisz_stan()
        index = current_row
        
        # Zamień elementy w liście obrazów
        self.ilustracje[index], self.ilustracje[index-1] = self.ilustracje[index-1], self.ilustracje[index]
        
        # Zamień opisy
        if index in self.opisy_krokow and index-1 in self.opisy_krokow:
            self.opisy_krokow[index], self.opisy_krokow[index-1] = self.opisy_krokow[index-1], self.opisy_krokow[index]
        
        # Ponownie załaduj listę z nowymi numerami kroków
        self.odswiez_liste()
        self.steps_list.setCurrentRow(index-1)
        self.pokaz_podglad()
    
    def przenies_w_dol(self):
        """Przenosi zaznaczony krok w dół"""
        current_row = self.steps_list.currentRow()
        if current_row < 0 or current_row >= len(self.ilustracje) - 1:
            return
        
        self.zapisz_stan()
        index = current_row
        
        # Zamień elementy w liście obrazów
        self.ilustracje[index], self.ilustracje[index+1] = self.ilustracje[index+1], self.ilustracje[index]
        
        # Zamień opisy
        if index in self.opisy_krokow and index+1 in self.opisy_krokow:
            self.opisy_krokow[index], self.opisy_krokow[index+1] = self.opisy_krokow[index+1], self.opisy_krokow[index]
        
        # Ponownie załaduj listę z nowymi numerami kroków
        self.odswiez_liste()
        self.steps_list.setCurrentRow(index+1)
        self.pokaz_podglad()
    
    def odswiez_liste(self):
        """Odświeża listę kroków z miniaturami i numerami kroków"""
        self.steps_list.clear()
        for i, sciezka in enumerate(self.ilustracje):
            step_number = i + 1
            if i in self.opisy_krokow and self.opisy_krokow[i]['nazwa']:
                nazwa_kroku = self.opisy_krokow[i]['nazwa']
            else:
                nazwa_pliku = os.path.basename(sciezka)
                nazwa_kroku = os.path.splitext(nazwa_pliku)[0]
            
            self.steps_list.add_image_item(sciezka, step_number, nazwa_kroku)
    
    def pokaz_podglad(self, item=None):
        """Pokazuje podgląd zaznaczonego kroku"""
        current_row = self.steps_list.currentRow()
        if current_row >= 0 and current_row < len(self.ilustracje):
            self.aktualny_wybrany_krok = current_row
            sciezka_obrazu = self.ilustracje[current_row]
            
            # Ustaw obraz w podglądzie
            self.preview_widget.set_image(sciezka_obrazu)
            self.preview_widget.current_image_path = sciezka_obrazu
            
            # Ustaw nazwę i opis w edytorze
            if current_row in self.opisy_krokow:
                self.step_name_edit.setText(self.opisy_krokow[current_row]['nazwa'])
                self.step_desc_edit.setPlainText(self.opisy_krokow[current_row]['opis'])
            else:
                self.step_name_edit.setText(f"Krok {current_row + 1}")
                self.step_desc_edit.setPlainText(self.t("domyslny_szablon"))
            
            self.aktualizuj_info_o_kroku()
    
    def aktualizuj_info_o_kroku(self):
        """Aktualizuje informację o aktualnie edytowanym kroku"""
        if self.aktualny_wybrany_krok is not None:
            nazwa_kroku = self.step_name_edit.text()
            if not nazwa_kroku.strip():
                if self.aktualny_wybrany_krok in self.opisy_krokow and self.opisy_krokow[self.aktualny_wybrany_krok]['nazwa']:
                    nazwa_kroku = self.opisy_krokow[self.aktualny_wybrany_krok]['nazwa']
                else:
                    nazwa_kroku = f"Krok {self.aktualny_wybrany_krok + 1}"
            
            self.step_info_label.setText(
                f"Edytujesz krok {self.aktualny_wybrany_krok + 1}: {nazwa_kroku}"
            )
            self.step_info_label.setStyleSheet("color: green; font-weight: bold; font-size: 12px;")
        else:
            self.step_info_label.setText(self.t("brak_zaznaczenia"))
            self.step_info_label.setStyleSheet("color: red; font-weight: bold; font-size: 12px;")
    
    def aktualizuj_nazwe_kroku_na_liscie(self):
        """Aktualizuje nazwę kroku na liście w czasie rzeczywistym"""
        if self.aktualny_wybrany_krok is not None:
            nowa_nazwa = self.step_name_edit.text()
            if nowa_nazwa.strip():
                # Aktualizuj listę z numerem kroku
                step_number = self.aktualny_wybrany_krok + 1
                if 0 <= self.aktualny_wybrany_krok < self.steps_list.count():
                    sciezka = self.ilustracje[self.aktualny_wybrany_krok]
                    self.steps_list.update_item(self.aktualny_wybrany_krok, sciezka, step_number, nowa_nazwa)
                
                self.aktualizuj_info_o_kroku()
    
    def zapisz_opis_kroku(self):
        """Zapisuje opis bieżącego kroku"""
        if self.aktualny_wybrany_krok is None:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("najpierw_wybierz_krok"))
            return
        
        self.zapisz_stan()
        index = self.aktualny_wybrany_krok
        nazwa = self.step_name_edit.text().strip()
        opis = self.step_desc_edit.toPlainText().strip()
        
        # Jeśli nazwa jest pusta, użyj domyślnej
        if not nazwa:
            nazwa = f"Krok {index + 1}"
        
        # Zapisz nazwę i opis
        self.opisy_krokow[index] = {
            'nazwa': nazwa,
            'opis': opis
        }
        
        # Aktualizuj listę z numerem kroku
        step_number = index + 1
        if 0 <= index < self.steps_list.count():
            sciezka = self.ilustracje[index]
            self.steps_list.update_item(index, sciezka, step_number, nazwa)
        
        self.aktualizuj_info_o_kroku()
        QMessageBox.information(self, self.t("sukces"), 
                              self.t("opis_kroku_zostal_zapisany"))
    
    def aktualizuj_podglad_opisu_na_zywo(self):
        """Aktualizuje podgląd opisu na żywo podczas edycji"""
        self.preview_timer.stop()
    
    def generuj_i_otworz_dokument(self):
        """Generuje dokumentację i automatycznie ją otwiera"""
        if not self.ilustracje:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("dodaj_przynajmniej_jedna_ilustracje"))
            return
        
        if not self.kod_edit.text() or not self.nazwa_edit.text():
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("wprowadz_kod_i_nazwe_dokumentu"))
            return
        
        # Zapytaj użytkownika gdzie zapisać dokument
        sciezka, _ = QFileDialog.getSaveFileName(
            self,
            self.t("wybierz_katalog"),
            "dokumentacja_kpk.docx",
            "Dokumenty Word (*.docx);;Wszystkie pliki (*.*)"
        )
        
        if not sciezka:  # Użytkownik anulował
            return
        
        self.sciezka_doc = sciezka
        
        try:
            # Pokaż dialog postępu
            progress = QProgressDialog(
                self.t("generowanie_dokumentu"),
                "Anuluj",
                0,
                len(self.ilustracje) + 5,
                self
            )
            progress.setWindowTitle("Generowanie dokumentu")
            progress.setWindowModality(Qt.WindowModal)
            progress.setValue(0)
            progress.show()
            
            doc = Document()
            progress.setValue(1)
            
            # Ustaw styl czcionki
            style = doc.styles['Normal']
            font = style.font
            font.name = self.font_combo.currentText()
            font.size = Pt(int(self.font_size_combo.currentText()))
            progress.setValue(2)
            
            # Tworzymy nagłówek i stopkę
            self.create_header_footer(doc)
            progress.setValue(3)
            
            # Strona tytułowa
            title_section = doc.add_heading(self.t("dokument_tytul"), 0)
            title_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            kod_label = self.t("kod_produktu").replace(":", "")
            nazwa_label = self.t("nazwa_dokumentu").replace(":", "")
            data_label = self.t("data_dokumentu").replace(":", "")
            autor_label = self.t("autor_dokumentu").replace(":", "")
            
            doc.add_paragraph().add_run(f"{kod_label}: {self.kod_edit.text()}").bold = True
            doc.add_paragraph().add_run(f"{nazwa_label}: {self.nazwa_edit.text()}").bold = True
            doc.add_paragraph().add_run(f"{data_label}: {self.data_edit.text()}").bold = True
            if self.autor_edit.text().strip():
                doc.add_paragraph().add_run(f"{autor_label}: {self.autor_edit.text()}").bold = True
            
            liczba_krokow_label = self.t("liczba_krokow").replace(":", "")
            doc.add_paragraph().add_run(f"{liczba_krokow_label} {len(self.ilustracje)}").bold = True
            
            doc.add_page_break()
            progress.setValue(4)
            
            # Spis treści
            doc.add_heading(self.t("spis_tresci"), level=1)
            for i in range(len(self.ilustracje)):
                nazwa_kroku = self.opisy_krokow.get(i, {}).get('nazwa', f'Krok {i+1}')
                doc.add_paragraph(f"Krok {i+1}: {nazwa_kroku}", style='List Number')
            
            doc.add_page_break()
            progress.setValue(5)
            
            # Procedura krok po kroku
            doc.add_heading(self.t("procedura"), level=1)
            
            # Dodaj każdą ilustrację z tekstem
            for i, sciezka_ilustracji in enumerate(self.ilustracje, 1):
                self.dodaj_krok_do_dokumentu(doc, sciezka_ilustracji, i)
                progress.setValue(5 + i)
                if progress.wasCanceled():
                    return
            
            # Strona końcowa
            doc.add_page_break()
            doc.add_heading(self.t("koniec"), level=1)
            doc.add_paragraph(self.t("dokumentacja_wygenerowana"))
            doc.add_paragraph(f"{self.t('data_generacji')} {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Zapisz dokument
            doc.save(self.sciezka_doc)
            progress.setValue(len(self.ilustracje) + 5)
            
            progress.close()
            
            # Otwórz dokument automatycznie
            if os.path.exists(self.sciezka_doc):
                try:
                    if sys.platform == "win32":
                        os.startfile(self.sciezka_doc)
                    elif sys.platform == "darwin":
                        subprocess.call(["open", self.sciezka_doc])
                    else:
                        subprocess.call(["xdg-open", self.sciezka_doc])
                    
                    QMessageBox.information(self, self.t("sukces"), 
                                          f"{self.t('dokumentacja_wygenerowana_i_otwarta')}:\n{self.sciezka_doc}")
                except:
                    # Jeśli nie udało się otworzyć, pokaż komunikat z lokalizacją
                    QMessageBox.information(self, self.t("sukces"), 
                                          f"{self.t('dokument_wygenerowany_ale_nie_otwarty')}:\n{self.sciezka_doc}")
            else:
                QMessageBox.warning(self, self.t("ostrzeżenie"), 
                                  self.t("dokument_wygenerowany_ale_nie_otwarty"))
            
        except Exception as e:
            error_details = traceback.format_exc()
            QMessageBox.critical(self, self.t("błąd"), 
                               f"{self.t('wystąpił_błąd_podczas_generowania')}:\n{str(e)}\n\nSzczegóły:\n{error_details}")
    
    def edytuj_w_paincie(self):
        """Otwiera zaznaczony obraz w programie Paint i zapisuje zmiany"""
        current_row = self.steps_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("najpierw_wybierz_obraz"))
            return
        
        index = current_row
        
        # SPRAWDŹ CZY OBRAZ ISTNIEJE
        if index >= len(self.ilustracje):
            QMessageBox.warning(self, self.t("błąd"), 
                              "Indeks poza zakresem listy obrazów!")
            return
        
        sciezka_obrazu = self.ilustracje[index]
        
        if not os.path.exists(sciezka_obrazu):
            QMessageBox.warning(self, self.t("błąd"), 
                              f"Obraz nie istnieje:\n{sciezka_obrazu}\n\nŚcieżka może być nieprawidłowa lub plik został usunięty.")
            return
        
        try:
            if sys.platform == "win32":
                # Utwórz KOPIĘ obrazu w stałym miejscu przed edycją
                # Zapisz w głównym katalogu tymczasowym z czytelną nazwą
                timestamp = int(datetime.now().timestamp() * 1000)
                nazwa_pliku = os.path.basename(sciezka_obrazu)
                nazwa_bez_ext = os.path.splitext(nazwa_pliku)[0]
                
                # Stwórz dedykowany katalog dla edytowanych obrazów
                edit_temp_dir = os.path.join(self.temp_dir, "paint_edit")
                os.makedirs(edit_temp_dir, exist_ok=True)
                
                # Utwórz unikalną ścieżkę dla kopii
                kopia_obrazu = os.path.join(
                    edit_temp_dir, 
                    f"edit_{index}_{timestamp}_{nazwa_bez_ext}.png"
                )
                
                # Skopiuj oryginalny obraz do lokalizacji edycji
                try:
                    with Image.open(sciezka_obrazu) as img:
                        # Zawsze zapisuj jako PNG dla lepszej jakości i kompresji bezstratnej
                        if img.mode in ('RGBA', 'LA', 'P'):
                            img.save(kopia_obrazu, 'PNG')
                        else:
                            img.convert('RGB').save(kopia_obrazu, 'PNG', quality=100)
                except Exception as e:
                    # Jeśli nie udało się przekonwertować, po prostu skopiuj
                    shutil.copy2(sciezka_obrazu, kopia_obrazu)
                
                # Upewnij się, że plik został utworzony
                if not os.path.exists(kopia_obrazu):
                    QMessageBox.warning(self, self.t("błąd"), 
                                      "Nie udało się utworzyć kopii obrazu do edycji!")
                    return
                
                # Użyj absolutnej ścieżki
                paint_path = "mspaint.exe"
                
                # Otwórz Paint z KOPIĄ obrazu
                try:
                    # Sprawdź czy Paint istnieje
                    process = subprocess.Popen([paint_path, kopia_obrazu])
                    
                    result = QMessageBox.question(
                        self,
                        self.t("edytowanie_w_paincie"),
                        self.t("czy_zapisac_zmiany"),
                        QMessageBox.Yes | QMessageBox.No
                    )
                    
                    if result == QMessageBox.Yes:
                        process.wait()
                        
                        # Sprawdź czy edytowany plik istnieje i ma rozsądny rozmiar
                        if os.path.exists(kopia_obrazu) and os.path.getsize(kopia_obrazu) > 100:
                            self.zapisz_stan()
                            
                            # Przenieś edytowany obraz z powrotem do oryginalnej ścieżki
                            try:
                                # Zapisz jako JPG dla zgodności z resztą systemu
                                final_sciezka = sciezka_obrazu
                                
                                # Jeśli oryginał był w katalogu tymczasowym, nadpisz
                                if self.temp_dir in sciezka_obrazu:
                                    with Image.open(kopia_obrazu) as edited_img:
                                        if edited_img.mode in ('RGBA', 'LA', 'P'):
                                            edited_img.convert('RGB').save(final_sciezka, 'JPEG', quality=90, optimize=True)
                                        else:
                                            edited_img.save(final_sciezka, 'JPEG', quality=90, optimize=True)
                                else:
                                    # Jeśli to oryginalny plik użytkownika, utwórz kopię z datą
                                    dir_path = os.path.dirname(sciezka_obrazu)
                                    nazwa_pliku = os.path.basename(sciezka_obrazu)
                                    nazwa_bez_ext = os.path.splitext(nazwa_pliku)[0]
                                    ext = os.path.splitext(sciezka_obrazu)[1]
                                    
                                    # Zachowaj kopię oryginału
                                    backup_sciezka = os.path.join(
                                        dir_path, 
                                        f"{nazwa_bez_ext}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
                                    )
                                    shutil.copy2(sciezka_obrazu, backup_sciezka)
                                    
                                    # Nadpisz oryginalny plik edytowanym
                                    with Image.open(kopia_obrazu) as edited_img:
                                        if edited_img.mode in ('RGBA', 'LA', 'P'):
                                            edited_img.convert('RGB').save(final_sciezka, 'JPEG', quality=90, optimize=True)
                                        else:
                                            edited_img.save(final_sciezka, 'JPEG', quality=90, optimize=True)
                                
                                # Zaktualizuj ścieżkę w liście
                                self.ilustracje[index] = final_sciezka
                                
                                # Odśwież miniaturę i podgląd
                                self.steps_list.setCurrentRow(index)
                                self.pokaz_podglad()
                                
                                QMessageBox.information(self, self.t("sukces"), 
                                                      self.t("obraz_zaktualizowany"))
                            except Exception as e:
                                QMessageBox.warning(self, self.t("błąd"), 
                                                  f"Błąd podczas zapisywania edytowanego obrazu: {str(e)}")
                        else:
                            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                                              "Edytowany plik jest pusty lub nie istnieje!")
                except FileNotFoundError:
                    QMessageBox.warning(self, self.t("błąd"), 
                                      "Nie znaleziono programu Paint (mspaint.exe).\nUpewnij się, że jest zainstalowany w systemie.")
                except Exception as e:
                    QMessageBox.warning(self, self.t("błąd"), 
                                      f"Błąd otwierania Paint:\n{str(e)}")
                
                # Sprzątanie - usuń kopię tymczasową
                try:
                    if os.path.exists(kopia_obrazu):
                        os.remove(kopia_obrazu)
                except:
                    pass
                    
            else:
                QMessageBox.warning(
                    self,
                    self.t("ostrzeżenie"),
                    self.t("funkcja_dostepna_tylko_windows")
                )
                
        except Exception as e:
            QMessageBox.critical(self, self.t("błąd"), 
                               f"{self.t('błąd_otwierania_paint')}: {e}")
            print(f"Błąd w edytuj_w_paincie: {traceback.format_exc()}")
    
    def create_element(self, name):
        return OxmlElement(name)
    
    def create_header_footer(self, doc):
        """Tworzy nagłówek i stopkę dla dokumentu"""
        
        # Nagłówek
        section = doc.sections[0]
        header = section.header
        
        # Tworzymy tabelę w nagłówku z 4 kolumnami
        header_table = header.add_table(rows=2, cols=4, width=Inches(6.5))
        header_table.autofit = False
        
        # Ustawienie szerokości kolumn
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(2.0)
        header_table.columns[2].width = Inches(1.5)
        header_table.columns[3].width = Inches(1.5)
        
        # Pierwszy wiersz - główne informacje
        cell1 = header_table.cell(0, 0)
        kod_label = self.t("kod_produktu").replace(":", "")
        cell1.text = f"{kod_label}: {self.kod_edit.text()}"
        cell1.paragraphs[0].runs[0].bold = True
        
        cell2 = header_table.cell(0, 1)
        cell2.merge(header_table.cell(0, 2))
        cell2.text = self.nazwa_edit.text()
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraphs[0].runs[0].bold = True
        
        cell3 = header_table.cell(0, 3)
        data_label = self.t("data_dokumentu").replace(":", "")
        cell3.text = f"{data_label}: {self.data_edit.text()}"
        cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell3.paragraphs[0].runs[0].bold = True
        
        cell4 = header_table.cell(1, 0)
        cell4.merge(header_table.cell(1, 3))
        autor_label = self.t("autor_dokumentu").replace(":", "")
        if self.autor_edit.text().strip():
            cell4.text = f"{autor_label}: {self.autor_edit.text()}"
        else:
            cell4.text = f"{autor_label}: {self.t('nie_podano')}"
        cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell4.paragraphs[0].runs[0].italic = True
        
        # Stopka z numeracją stron
        footer = section.footer
        footer_table = footer.add_table(rows=1, cols=1, width=Inches(6.5))
        footer_table.autofit = False
        
        footer_cell = footer_table.cell(0, 0)
        footer_paragraph = footer_cell.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dodajemy numer strony
        run = footer_paragraph.add_run()
        fldChar = self.create_element('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = self.create_element('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = self.create_element('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run = footer_paragraph.add_run(" / ")
        
        run2 = footer_paragraph.add_run()
        fldChar3 = self.create_element('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = self.create_element('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        
        fldChar4 = self.create_element('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        
        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)
    
    def dodaj_krok_do_dokumentu(self, doc, sciezka_ilustracji, numer):
        """Dodaje krok do dokumentu z obrazem i opisem"""
        index = numer - 1
        nazwa_kroku = self.opisy_krokow.get(index, {}).get('nazwa', f'Krok {numer}')
        opis_kroku = self.opisy_krokow.get(index, {}).get('opis', self.t("domyslny_szablon"))
        
        # Nagłówek kroku
        doc.add_heading(f'Krok {numer}: {nazwa_kroku}', level=2)
        
        # Konwersja rozmiaru na cale
        try:
            rozmiar_cm = float(self.size_combo.currentText())
            rozmiar_cale = rozmiar_cm / 2.54
        except:
            rozmiar_cale = 3.15  # Domyślnie 8 cm
        
        układ = self.layout_combo.currentData()
        
        if układ == "lewo_prawo":
            tabela = doc.add_table(rows=1, cols=2)
            tabela.autofit = False
            
            szerokosc_strony = 6.5
            szerokosc_obrazu = rozmiar_cale
            szerokosc_tekstu = szerokosc_strony - szerokosc_obrazu - 0.5
            
            tabela.columns[0].width = Inches(szerokosc_obrazu)
            tabela.columns[1].width = Inches(szerokosc_tekstu)
            
            komorka_obraz = tabela.cell(0, 0)
            komorka_tekst = tabela.cell(0, 1)
            
            # Obraz po lewej
            akapit_obraz = komorka_obraz.paragraphs[0]
            akapit_obraz.alignment = WD_ALIGN_PARAGRAPH.CENTER
            uruchomienie_obraz = akapit_obraz.add_run()
            
            # Sprawdź czy obraz istnieje
            if os.path.exists(sciezka_ilustracji):
                uruchomienie_obraz.add_picture(sciezka_ilustracji, width=Inches(rozmiar_cale))
            else:
                # Jeśli obraz nie istnieje, dodaj tekst informacyjny
                uruchomienie_obraz.add_text(f"[Brak obrazu dla kroku {numer}]")
            
            # Tekst instrukcji po prawej
            akapit_tekst = komorka_tekst.paragraphs[0]
            akapit_tekst.add_run(f"{self.t('instrukcja_wykonania')}\n").bold = True
            self.dodaj_sformatowany_tekst(doc, akapit_tekst, opis_kroku)
            
        elif układ == "gora":
            # Obraz na górze
            akapit_obraz = doc.add_paragraph()
            akapit_obraz.alignment = WD_ALIGN_PARAGRAPH.CENTER
            uruchomienie_obraz = akapit_obraz.add_run()
            
            if os.path.exists(sciezka_ilustracji):
                uruchomienie_obraz.add_picture(sciezka_ilustracji, width=Inches(rozmiar_cale))
            else:
                uruchomienie_obraz.add_text(f"[Brak obrazu dla kroku {numer}]")
            
            doc.add_paragraph(f"{self.t('instrukcja_wykonania')}\n").bold = True
            self.dodaj_sformatowany_tekst_po_akapicie(doc, opis_kroku)
            
        else:  # układ == "dol"
            # Obraz na dole
            doc.add_paragraph(f"{self.t('instrukcja_wykonania')}\n").bold = True
            self.dodaj_sformatowany_tekst_po_akapicie(doc, opis_kroku)
            
            akapit_obraz = doc.add_paragraph()
            akapit_obraz.alignment = WD_ALIGN_PARAGRAPH.CENTER
            uruchomienie_obraz = akapit_obraz.add_run()
            
            if os.path.exists(sciezka_ilustracji):
                uruchomienie_obraz.add_picture(sciezka_ilustracji, width=Inches(rozmiar_cale))
            else:
                uruchomienie_obraz.add_text(f"[Brak obrazu dla kroku {numer}]")
        
        # Dodaj odstęp przed następnym krokiem
        doc.add_paragraph()
    
    def dodaj_sformatowany_tekst(self, doc, akapit, tekst):
        """Dodaje sformatowany tekst z obsługą podpunktów i sekcji Maszyna"""
        linie = tekst.split('\n')
        for linia in linie:
            linia = linia.strip()
            if not linia:
                akapit.add_run('\n')
            elif linia.startswith('Maszyna:') or linia.startswith('Machine:'):
                akapit.add_run('\n' + linia + '\n').bold = True
            elif linia.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and '.' in linia:
                akapit.add_run('\n' + linia + '\n').bold = True
            elif linia.startswith('•'):
                akapit.add_run(linia + '\n')
            else:
                akapit.add_run(linia + '\n')
    
    def dodaj_sformatowany_tekst_po_akapicie(self, doc, tekst):
        """Dodaje sformatowany tekst jako osobne akapity"""
        linie = tekst.split('\n')
        for linia in linie:
            linia = linia.strip()
            if not linia:
                doc.add_paragraph()
            elif linia.startswith('Maszyna:') or linia.startswith('Machine:'):
                p = doc.add_paragraph()
                p.add_run(linia).bold = True
            elif linia.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and '.' in linia:
                p = doc.add_paragraph()
                p.add_run(linia).bold = True
            elif linia.startswith('•'):
                p = doc.add_paragraph()
                p.add_run("    " + linia)
            else:
                doc.add_paragraph(linia)
    
    def zapisz_projekt(self):
        """Zapisuje cały projekt do pliku JSON z osadzonymi obrazami"""
        if not self.ilustracje:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              self.t("brak_danych_do_zapisania"))
            return
        
        sciezka, _ = QFileDialog.getSaveFileName(
            self,
            "Zapisz projekt",
            "",
            "Pliki projektu (*.json)"
        )
        
        if sciezka:
            try:
                obrazy_base64 = {}
                for i, sciezka_img in enumerate(self.ilustracje):
                    try:
                        with open(sciezka_img, 'rb') as f:
                            obraz_data = f.read()
                            obraz_base64 = base64.b64encode(obraz_data).decode('utf-8')
                            nazwa_pliku = os.path.basename(sciezka_img)
                            obrazy_base64[str(i)] = {
                                'nazwa': nazwa_pliku,
                                'dane': obraz_base64
                            }
                    except Exception as e:
                        QMessageBox.warning(self, self.t("ostrzeżenie"), 
                                          f"{self.t('błąd_odczytu_obrazu')} {os.path.basename(sciezka_img)}: {e}")
                
                opisy_krokow_str_keys = {}
                for key, value in self.opisy_krokow.items():
                    opisy_krokow_str_keys[str(key)] = value
                
                projekt = {
                    'kod': self.kod_edit.text(),
                    'nazwa': self.nazwa_edit.text(),
                    'data': self.data_edit.text(),
                    'autor': self.autor_edit.text(),
                    'obrazy': obrazy_base64,
                    'opisy_krokow': opisy_krokow_str_keys,
                    'ustawienia': {
                        'uklad': self.layout_combo.currentData(),
                        'rozmiar': self.size_combo.currentText(),
                        'czcionka': self.font_combo.currentText(),
                        'rozmiar_czcionki': self.font_size_combo.currentText(),
                        'jezyk': self.jezyk,
                        'dark_theme': self.dark_theme
                    },
                    'tlumaczenia': self.tlumaczenia
                }
                
                with open(sciezka, 'w', encoding='utf-8') as f:
                    json.dump(projekt, f, ensure_ascii=False, indent=2)
                
                QMessageBox.information(self, self.t("sukces"), 
                                      f"{self.t('projekt_zapisany')}: {sciezka}\n\n{self.t('obrazy_zostaly_osadzone')}")
                
            except Exception as e:
                QMessageBox.critical(self, self.t("błąd"), 
                                   f"{self.t('nie_udało_się_zapisać_projektu')}: {e}")
    
    def wczytaj_projekt(self):
        """Wczytuje projekt z pliku JSON i odtwarza obrazy"""
        sciezka, _ = QFileDialog.getOpenFileName(
            self,
            "Wczytaj projekt",
            "",
            "Pliki projektu (*.json)"
        )
        
        if sciezka:
            try:
                with open(sciezka, 'r', encoding='utf-8') as f:
                    projekt = json.load(f)
                
                self.kod_edit.setText(projekt.get('kod', ''))
                self.nazwa_edit.setText(projekt.get('nazwa', ''))
                self.data_edit.setText(projekt.get('data', ''))
                self.autor_edit.setText(projekt.get('autor', ''))
                
                self.ilustracje = []
                obrazy_base64 = projekt.get('obrazy', {})
                
                for i in range(len(obrazy_base64)):
                    str_i = str(i)
                    if str_i in obrazy_base64:
                        obraz_data = obrazy_base64[str_i]
                        nazwa_pliku = obraz_data['nazwa']
                        dane_base64 = obraz_data['dane']
                        
                        obraz_binary = base64.b64decode(dane_base64)
                        sciezka_obrazu = os.path.join(self.temp_dir, f"loaded_{i}_{nazwa_pliku}")
                        with open(sciezka_obrazu, 'wb') as f_img:
                            f_img.write(obraz_binary)
                        
                        self.ilustracje.append(sciezka_obrazu)
                
                opisy_krokow_int_keys = {}
                opisy_krokow_str_keys = projekt.get('opisy_krokow', {})
                for key, value in opisy_krokow_str_keys.items():
                    opisy_krokow_int_keys[int(key)] = value
                self.opisy_krokow = opisy_krokow_int_keys
                
                ustawienia = projekt.get('ustawienia', {})
                
                # Ustaw układ
                layout_value = ustawienia.get('uklad', 'lewo_prawo')
                index = self.layout_combo.findData(layout_value)
                if index >= 0:
                    self.layout_combo.setCurrentIndex(index)
                
                self.size_combo.setCurrentText(ustawienia.get('rozmiar', '8'))
                self.font_combo.setCurrentText(ustawienia.get('czcionka', 'Arial'))
                self.font_size_combo.setCurrentText(ustawienia.get('rozmiar_czcionki', '11'))
                
                self.jezyk = ustawienia.get('jezyk', 'polski')
                index = self.lang_combo.findData(self.jezyk)
                if index >= 0:
                    self.lang_combo.setCurrentIndex(index)
                
                # Ustaw motyw
                dark_theme = ustawienia.get('dark_theme', False)
                self.zmien_motyw(dark_theme)
                
                if 'tlumaczenia' in projekt:
                    self.tlumaczenia = projekt['tlumaczenia']
                
                self.odswiez_liste()
                self.odswiez_interfejs()
                
                self.stan_historia = []
                self.aktualny_stan_index = -1
                self.zapisz_stan()
                
                QMessageBox.information(self, self.t("sukces"), 
                                      f"{self.t('projekt_wczytany')}: {sciezka}\n\n{self.t('załadowano_obrazy')} {len(self.ilustracje)} {self.t('obrazów')}")
                
            except Exception as e:
                QMessageBox.critical(self, self.t("błąd"), 
                                   f"{self.t('nie_udało_się_wczytać_projektu')}: {e}")
    
    def wyczysc_wszystko(self):
        """Czyści wszystkie dane projektu"""
        self.zapisz_stan()
        self.ilustracje.clear()
        self.opisy_krokow.clear()
        self.steps_list.clear()
        self.preview_widget.set_image(None)
        self.step_name_edit.clear()
        self.step_desc_edit.clear()
        self.kod_edit.setText("xxx-xxxx-xxx")
        self.nazwa_edit.setText("Dokumentacja produktu")
        self.data_edit.setText(datetime.now().strftime("%Y-%m-%d"))
        self.autor_edit.setText("Dorota Zaręba")
        self.aktualny_wybrany_krok = None
        self.aktualizuj_info_o_kroku()
    
    def utworz_kopie_zapasowa(self):
        """Tworzy kopię zapasową wszystkich obrazów i danych projektu"""
        if not self.ilustracje:
            QMessageBox.warning(self, self.t("ostrzeżenie"), 
                              "Brak danych do kopii zapasowej!")
            return
        
        katalog = QFileDialog.getExistingDirectory(
            self,
            self.t("katalog_dla_kopii_zapasowej")
        )
        
        if not katalog:
            return
        
        try:
            # Utwórz podkatalog z datą
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_dir = os.path.join(katalog, f"backup_dokumentacji_{timestamp}")
            os.makedirs(backup_dir, exist_ok=True)
            
            # Skopiuj wszystkie obrazy
            for i, sciezka in enumerate(self.ilustracje):
                if os.path.exists(sciezka):
                    nazwa_pliku = os.path.basename(sciezka)
                    dest_path = os.path.join(backup_dir, f"krok_{i+1}_{nazwa_pliku}")
                    shutil.copy2(sciezka, dest_path)
            
            # Zapisz dane projektu
            dane_projektu = {
                'kod': self.kod_edit.text(),
                'nazwa': self.nazwa_edit.text(),
                'data': self.data_edit.text(),
                'autor': self.autor_edit.text(),
                'opisy_krokow': self.opisy_krokow,
                'data_utworzenia_kopii': datetime.now().isoformat()
            }
            
            with open(os.path.join(backup_dir, "dane_projektu.json"), 'w', encoding='utf-8') as f:
                json.dump(dane_projektu, f, ensure_ascii=False, indent=2)
            
            QMessageBox.information(self, self.t("sukces"), 
                                  f"{self.t('kopia_zapasowa_utworzona')}:\n{backup_dir}")
            
        except Exception as e:
            QMessageBox.critical(self, self.t("błąd"), 
                               f"Błąd tworzenia kopii zapasowej: {e}")
    
    def autozapisz_projekt(self):
        """Automatycznie zapisuje projekt co określony interwał"""
        if not self.ilustracje or not self.autosave_enabled:
            return
        
        try:
            # Generuj nazwę pliku
            kod = self.kod_edit.text().strip()
            if not kod or kod == "xxx-xxxx-xxx":
                kod = "bez_kodu"
            else:
                # Usuń niebezpieczne znaki z nazwy pliku
                kod = "".join(c for c in kod if c.isalnum() or c in ('-', '_'))
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nazwa_pliku = f"{kod}_{timestamp}.json"
            sciezka = os.path.join(self.autosave_dir, nazwa_pliku)
            
            # Przygotuj dane do zapisu
            obrazy_base64 = {}
            for i, sciezka_img in enumerate(self.ilustracje):
                try:
                    if os.path.exists(sciezka_img):
                        with open(sciezka_img, 'rb') as f:
                            obraz_data = f.read()
                            obraz_base64 = base64.b64encode(obraz_data).decode('utf-8')
                            nazwa_pliku_img = os.path.basename(sciezka_img)
                            obrazy_base64[str(i)] = {
                                'nazwa': nazwa_pliku_img,
                                'dane': obraz_base64
                            }
                except:
                    continue  # Pomiń obrazy, które nie mogą być odczytane
            
            opisy_krokow_str_keys = {}
            for key, value in self.opisy_krokow.items():
                opisy_krokow_str_keys[str(key)] = value
            
            projekt = {
                'kod': self.kod_edit.text(),
                'nazwa': self.nazwa_edit.text(),
                'data': self.data_edit.text(),
                'autor': self.autor_edit.text(),
                'obrazy': obrazy_base64,
                'opisy_krokow': opisy_krokow_str_keys,
                'ustawienia': {
                    'uklad': self.layout_combo.currentData(),
                    'rozmiar': self.size_combo.currentText(),
                    'czcionka': self.font_combo.currentText(),
                    'rozmiar_czcionki': self.font_size_combo.currentText(),
                    'jezyk': self.jezyk,
                    'dark_theme': self.dark_theme
                },
                'tlumaczenia': self.tlumaczenia,
                'timestamp_autosave': datetime.now().isoformat()
            }
            
            # Zapisz plik
            with open(sciezka, 'w', encoding='utf-8') as f:
                json.dump(projekt, f, ensure_ascii=False, indent=2)
            
            # Ogranicz liczbę plików autozapisu do 10 najnowszych
            self.oczysc_stare_autozapisy()
            
            # Logowanie (opcjonalne)
            print(f"Autozapisano projekt: {sciezka}")
            
        except Exception as e:
            print(f"Błąd autozapisu: {e}")
    
    def oczysc_stare_autozapisy(self):
        """Usuwa stare pliki autozapisu, zostawiając tylko 10 najnowszych"""
        try:
            # Znajdź wszystkie pliki autozapisu
            pliki = []
            for plik in os.listdir(self.autosave_dir):
                if plik.endswith('.json') and (plik.startswith(('bez_kodu_', 'xxx-xxxx-xxx_')) or '_' in plik):
                    sciezka = os.path.join(self.autosave_dir, plik)
                    if os.path.isfile(sciezka):
                        pliki.append((sciezka, os.path.getmtime(sciezka)))
            
            # Posortuj od najstarszego do najnowszego
            pliki.sort(key=lambda x: x[1])
            
            # Usuń najstarsze, zostawiając tylko 10 najnowszych
            while len(pliki) > 10:
                stary_plik, _ = pliki.pop(0)
                try:
                    os.remove(stary_plik)
                    print(f"Usunięto stary autozapis: {stary_plik}")
                except:
                    pass
                    
        except Exception as e:
            print(f"Błąd czyszczenia starych autozapisów: {e}")


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    
    # Domyślny jasny motyw
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(240, 240, 240))
    palette.setColor(QPalette.WindowText, Qt.black)
    palette.setColor(QPalette.Base, QColor(255, 255, 255))
    palette.setColor(QPalette.AlternateBase, QColor(240, 240, 240))
    palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
    palette.setColor(QPalette.ToolTipText, Qt.black)
    palette.setColor(QPalette.Text, Qt.black)
    palette.setColor(QPalette.Button, QColor(240, 240, 240))
    palette.setColor(QPalette.ButtonText, Qt.black)
    palette.setColor(QPalette.BrightText, Qt.red)
    palette.setColor(QPalette.Link, QColor(42, 130, 218))
    palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
    palette.setColor(QPalette.HighlightedText, Qt.white)
    app.setPalette(palette)
    
    window = GeneratorDokumentow()
    window.show()
    
    sys.exit(app.exec())


if __name__ == "__main__":
    main()